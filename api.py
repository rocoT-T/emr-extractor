from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from docx import Document
import re
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, Date, Text, DateTime
from sqlalchemy.orm import declarative_base, sessionmaker
import os
import tempfile


from dotenv import load_dotenv
load_dotenv()   # 放在所有配置之前，这样后面 os.getenv 就能读到 .env 里的值

# --------------------------
# 全局配置开关（开发/生产环境切换）
# --------------------------
DEBUG = os.getenv("DEBUG", "False").lower() == "true"  # 开发阶段设为True，显示详细错误；上线前改为False

# --------------------------
# 数据库配置
# --------------------------
# 从环境变量读取数据库连接地址
DATABASE_URL = os.getenv("DATABASE_URL")

# 如果环境变量没设置，给一个默认值（方便本地开发）
# if not DATABASE_URL:
#     DATABASE_URL = "mysql+pymysql://root:你的密码@localhost:3306/emr_db?charset=utf8mb4"

# 创建数据库引擎（echo=True 会打印 SQL 语句，调试时有用）
engine = create_engine(DATABASE_URL, echo=True if DEBUG else False)

# 会话工厂
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# 模型基类
Base = declarative_base()

# 定义数据库表模型（对应你建的 emr_records 表）
class EMRRecord(Base):
    __tablename__ = "emr_records"

    id = Column(Integer, primary_key=True, autoincrement=True)
    patient_name = Column(String(50))
    gender = Column(String(10))
    age = Column(Integer)
    hospital_id = Column(String(50))
    department = Column(String(50))
    bed_number = Column(String(20))
    admission_date = Column(Date)
    discharge_date = Column(Date)
    diagnosis = Column(Text)
    medical_order = Column(Text)
    filename = Column(String(255))
    created_at = Column(DateTime, default=datetime.now)

# 创建所有表（如果表已存在则跳过）
Base.metadata.create_all(bind=engine)

# 保存到数据库的函数
def save_to_db(patient_info: dict, filename: str):
    """将提取的患者信息保存到数据库，返回新记录的 ID"""
    db = SessionLocal()
    try:
        # 映射中文字段名到数据库列名
        record = EMRRecord(
            patient_name=patient_info.get("姓名"),
            gender=patient_info.get("性别"),
            age=patient_info.get("年龄"),
            hospital_id=patient_info.get("住院号"),
            department=patient_info.get("科室"),
            bed_number=patient_info.get("床号"),
            admission_date=parse_date(patient_info.get("入院日期")),  # 转为 datetime 对象
            discharge_date=parse_date(patient_info.get("出院日期")),
            diagnosis=patient_info.get("诊断"),
            medical_order=patient_info.get("医嘱"),
            filename=filename
        )
        db.add(record)
        db.commit()
        db.refresh(record)  # 获取数据库自动生成的 id
        return record.id
    except Exception as e:
        db.rollback()
        raise e
    finally:
        db.close()

# --------------------------
# (api初期)复制 main.py 里的所有函数过来
# --------------------------
def read_docx(file_path, verbose=False):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    for table in doc.tables:
        for row in table.rows:
            # 过滤掉全空的单元格，避免生成连续的制表符
            # 原代码：row_cells = [cell.text.strip() for cell in row.cells]
            # 问题：空单元格strip()后是空串，join会产生"\t\t\t"这样的垃圾数据
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            row_str = "\t".join(row_cells)
            # 只有当行不为空时才加入全文，避免空行
            if row_str:
                full_text.append(row_str)
    return "\n".join(full_text), doc

def normalize_date_str(date_str):
    """
    日期字符串标准化预处理
    将所有格式统一为 YYYY-MM-DD 格式，自动补全前导零
    支持：2025-5-19、2025/5/19、2025.5.19、2025年5月19日
    """
    if not date_str:
        return None
    
    # 统一所有分隔符为横杠
    normalized = re.sub(r'[/.年]', '-', date_str)
    # 移除"月"和"日"字
    normalized = normalized.replace('月', '-').replace('日', '')
    
    # 拆分年月日并补零
    parts = normalized.split('-')
    if len(parts) == 3:
        year, month, day = parts
        # 月份补零
        if len(month) == 1:
            month = '0' + month
        # 日期补零
        if len(day) == 1:
            day = '0' + day
        return f"{year}-{month}-{day}"
    
    return date_str

def parse_date(date_str):
    """
    智能解析日期字符串，支持多种格式
    返回datetime对象，解析失败返回None
    """
    if not date_str:
        return None
    
    # 先标准化日期字符串
    normalized_date = normalize_date_str(date_str)
    
    # 只需要保留标准格式即可
    formats = ["%Y-%m-%d"]
    for fmt in formats:
        try:
            return datetime.strptime(normalized_date, fmt)
        except ValueError:
            continue
    
    return None

def extract_from_tables(doc):
    result = {}
    field_mapping = {
        "姓名": ["姓名", "患者姓名", "病人姓名"],
        "性别": ["性别"],
        "年龄": ["年龄", "岁数"],
        "住院号": ["住院号", "病案号", "病历号", "住院编号"],
        "科室": ["科室", "入院科室", "所在科室"],
        "床号": ["床号", "床位号"],
        "入院日期": ["入院日期", "入院时间", "住院日期"],
        "出院日期": ["出院日期", "出院时间"]
    }
    for table in doc.tables:
        if len(table.rows) != 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        values = [cell.text.strip() for cell in table.rows[1].cells]
        for i, header in enumerate(headers):
            if i >= len(values):
                continue
            clean_header = header.replace(" ", "").replace("　", "").rstrip("：:")
            for standard_field, variants in field_mapping.items():
                if any(variant in clean_header for variant in variants):
                    result[standard_field] = values[i].strip()
                    break
    return result

def extract_from_fulltext(full_text):
    result = {}
    patterns = {
        "姓名": r"^(?:患者)?姓名[：:]\s*(.+)$",
        "性别": r"^性别[：:]\s*(.+)$",
        "年龄": r"^年龄[：:]\s*(.+)$",
        "住院号": r"^住院号[：:]\s*(.+)$",
        "入院日期": r"^入院日期[：:]\s*(.+)$",
        "出院日期": r"^出院日期[：:]\s*(.+)$",
        "诊断": r"^(?:入院|出院|主要)?诊断[：:]\s*(.+)$",
        "医嘱": r"^医嘱[：:]\s*(.+)$"
    }
    lines = full_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for field, pattern in patterns.items():
            match = re.match(pattern, line)
            if match:
                result[field] = match.group(1).strip()
                break
    return result

def clean_data(raw_data):
    cleaned = raw_data.copy()
    if cleaned.get("年龄") is not None:
        age_match = re.search(r"\d+", str(cleaned["年龄"]))
        if age_match:
            cleaned["年龄"] = int(age_match.group())
        else:
            cleaned["年龄"] = None
    admission_date = parse_date(cleaned.get("入院日期"))
    discharge_date = parse_date(cleaned.get("出院日期"))
    if admission_date and discharge_date and discharge_date < admission_date:
        cleaned["出院日期"] = None
    return cleaned

def extract_patient_info(full_text, doc):
    table_data = extract_from_tables(doc)
    text_data = extract_from_fulltext(full_text)
    raw_data = {
        "姓名": None,
        "性别": None,
        "年龄": None,
        "住院号": None,
        "科室": None,
        "床号": None,
        "入院日期": None,
        "出院日期": None,
        "诊断": None,
        "医嘱": None
    }
    for key in raw_data:
        raw_data[key] = table_data.get(key) if table_data.get(key) is not None else text_data.get(key)
    return clean_data(raw_data)

# --------------------------
# FastAPI 应用
# --------------------------
app = FastAPI(title="电子病历信息提取器API", version="0.2.1")

@app.get("/")
def root():
    return {
        "message": "EMR Extractor API is running",
        "version": "0.2.1",
        "status": "success",
        "debug_mode": DEBUG
    }

@app.post("/extract", summary="提取病历信息")
async def extract_emr(file: UploadFile = File(..., description="上传Word格式病历文件(.docx)")):
    # 1. 校验文件类型
    if not file.filename.endswith(".docx"):
        return JSONResponse(status_code=400, content={"error": "仅支持 .docx 文件"})

    # 2. 写入临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        # 3. 调用提取函数
        full_content, doc = read_docx(tmp_path)
        patient_info = extract_patient_info(full_content, doc)

        # 4. 保存到数据库
        try:
            record_id = save_to_db(patient_info, file.filename)
        except Exception as db_error:
            # 数据库保存失败，不影响接口返回，但会记录日志
            print(f"数据库保存失败: {db_error}")
            record_id = None


        # 5. 结果整理
        result = {}
        for k, v in patient_info.items():
            result[k] = v if v is not None else "未提取到"

        return {
            "filename": file.filename,
            "status": "success",
            "record_id": record_id,
            "data": result
        }

    except Exception as e:
        # 增加DEBUG开关，控制错误信息暴露程度
        if DEBUG:
            # 开发阶段：返回详细错误信息，方便调试
            return JSONResponse(status_code=500, content={"error": f"处理失败：{str(e)}"})
        else:
            # 生产阶段：只返回通用错误，避免泄露敏感信息
            return JSONResponse(status_code=500, content={"error": "服务器内部错误，请稍后重试"})

    finally:
        # 5. 删除临时文件
        os.unlink(tmp_path)