from docx import Document
import re
from datetime import datetime
from pathlib import Path
import pandas as pd

def read_docx(file_path, verbose=True):
    """
    读取Word文档
    :param verbose: 是否打印详细的段落和表格内容（单文件调试用True，批量处理用False）
    """
    doc = Document(file_path)
    full_text = []

    if verbose:
        print("===== 段落内容 =====")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if verbose:
                print(f"[段落] {text}")
            full_text.append(text)

    if verbose:
        print("\n===== 表格内容 =====")
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            row_str = "\t".join(row_cells)
            if verbose:
                print(f"[表格行] {row_str}")
            full_text.append(row_str)

    full_content = "\n".join(full_text)
    if verbose:
        print("\n===== 全文汇总 =====")
        print(full_content)

    return full_content, doc

def parse_date(date_str):
    """智能解析日期字符串，支持多种格式，解析失败返回None"""
    if not date_str:
        return None
    formats = ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return None

# --------------------------
# 分离关注点1：仅从表格提取（简化为只支持最常用的两行式）
# --------------------------
def extract_from_tables(doc):
    """从表格提取信息，仅支持最常见的两行式（表头在上，值在下）"""
    result = {}
    field_mapping = {
        "姓名": ["姓名", "患者姓名", "病人姓名"],
        "性别": ["性别"],
        "年龄": ["年龄", "岁数"],
        "住院号": ["住院号", "病案号", "病历号", "住院编号"],
        "科室": ["科室", "入院科室", "所在科室"],
        "床号": ["床号", "床位号"],
        "入院日期": ["入院日期", "入院时间", "住院日期"],
        "出院日期": ["出院日期", "出院时间"]
    }

    for table in doc.tables:
        if len(table.rows) != 2:
            continue  # 只处理两行式表格，其他格式忽略
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        values = [cell.text.strip() for cell in table.rows[1].cells]
        
        for i, header in enumerate(headers):
            if i >= len(values):
                continue
            clean_header = header.replace(" ", "").replace("　", "").rstrip("：:")
            for standard_field, variants in field_mapping.items():
                if any(variant in clean_header for variant in variants):
                    result[standard_field] = values[i].strip()
                    break
    return result

# --------------------------
# 分离关注点2：仅从纯文本提取（简化为逐行匹配，去掉复杂前向断言）
# --------------------------
def extract_from_fulltext(full_text):
    """从纯文本逐行提取信息，逻辑简单稳定，易调试"""
    result = {}
    patterns = {
        "姓名": r"^(?:患者)?姓名[：:]\s*(.+)$",
        "性别": r"^性别[：:]\s*(.+)$",
        "年龄": r"^年龄[：:]\s*(.+)$",
        "住院号": r"^住院号[：:]\s*(.+)$",
        "入院日期": r"^入院日期[：:]\s*(.+)$",
        "出院日期": r"^出院日期[：:]\s*(.+)$",
        "诊断": r"^(?:入院|出院|主要)?诊断[：:]\s*(.+)$",
        "医嘱": r"^医嘱[：:]\s*(.+)$"
    }

    lines = full_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for field, pattern in patterns.items():
            match = re.match(pattern, line)
            if match:
                result[field] = match.group(1).strip()
                break
    return result

# --------------------------
# 分离关注点3：统一数据清洗（仅做清洗，不做展示）
# --------------------------
def clean_data(raw_data):
    """统一清洗原始提取数据，返回清洗后的字典"""
    cleaned = raw_data.copy()

    # 修复：只有当年龄不为None时才清洗，支持0岁的情况
    if cleaned.get("年龄") is not None:
        age_match = re.search(r"\d+", str(cleaned["年龄"]))
        if age_match:
            cleaned["年龄"] = int(age_match.group())
        else:
            cleaned["年龄"] = None

    # 日期逻辑校验
    admission_date = parse_date(cleaned.get("入院日期"))
    discharge_date = parse_date(cleaned.get("出院日期"))
    if admission_date and discharge_date and discharge_date < admission_date:
        print(f"\n 警告：检测到出院日期({cleaned['出院日期']})早于入院日期({cleaned['入院日期']})，已自动置为未提取到")
        cleaned["出院日期"] = None

    return cleaned

# --------------------------
# 主提取函数：合并结果 + 清洗
# --------------------------
def extract_patient_info(full_text, doc):
    """主提取函数：表格优先，文本补全，最后统一清洗"""
    # 1. 分别提取
    table_data = extract_from_tables(doc)
    text_data = extract_from_fulltext(full_text)

    # 2. 合并结果（修复：只有表格值为None时才用文本值，避免空字符串覆盖）
    raw_data = {
        "姓名": None,
        "性别": None,
        "年龄": None,
        "住院号": None,
        "科室": None,
        "床号": None,
        "入院日期": None,
        "出院日期": None,
        "诊断": None,
        "医嘱": None
    }

    for key in raw_data:
        raw_data[key] = table_data.get(key) if table_data.get(key) is not None else text_data.get(key)

    # 3. 统一清洗
    return clean_data(raw_data)

# --------------------------
# 批量处理函数 + Excel导出
# --------------------------
def process_directory(directory_path):
    """批量处理目录下所有docx文件，导出结果到Excel"""
    results = []
    docx_files = list(Path(directory_path).glob("*.docx"))

    if not docx_files:
        print(f"错误：目录 {directory_path} 下没有找到docx文件")
        return

    print(f"\n开始批量处理，共找到 {len(docx_files)} 个病历文件\n")

    for file_path in docx_files:
        print(f"正在处理：{file_path.name}")
        try:
            # 修复：批量处理时关闭详细打印，只显示进度
            full_content, doc = read_docx(file_path, verbose=False)
            patient_info = extract_patient_info(full_content, doc)
            # 添加文件名作为标识
            patient_info["文件名"] = file_path.name
            results.append(patient_info)
            print(f" {file_path.name} 处理完成")
            print("-" * 50)
        except Exception as e:
            print(f" {file_path.name} 处理失败：{e}")
            print("-" * 50)
            continue

    # 导出到Excel
    if results:
        df = pd.DataFrame(results)
        # 调整列顺序，把文件名放第一列
        cols = ["文件名"] + [col for col in df.columns if col != "文件名"]
        df = df[cols]
        df.to_excel("output.xlsx", index=False)
        print(f"\n 批量处理完成！共成功处理 {len(results)} 个文件")
        print("结果已导出到：output.xlsx")

if __name__ == "__main__":
    print("请选择运行模式：")
    print("1. 处理单个文件（samples/patient_01.docx）")
    print("2. 批量处理samples目录下所有文件并导出Excel")
    choice = input("请输入数字（1/2）：")

    if choice == "1":
        file_path = "samples/patient_01.docx"
        try:
            # 单文件模式保留详细打印
            full_content, doc = read_docx(file_path, verbose=True)
            print("\n" + "="*50)
            print("\n===== 提取的患者信息 =====")
            patient_info = extract_patient_info(full_content, doc)
            # 展示层：统一处理未提取到的字段
            for key, value in patient_info.items():
                print(f"{key}: {value if value is not None else '未提取到'}")
        except Exception as e:
            print(f"处理失败：{e}")

    elif choice == "2":
        process_directory("samples")

    else:
        print("输入错误，请重新运行程序")