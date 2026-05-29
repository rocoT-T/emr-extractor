import random
import os
import json
from faker import Faker
from docx import Document
from datetime import timedelta

fake = Faker('zh_CN')

DEPARTMENTS = ['心内科', '呼吸科', '神经内科', '骨科', '消化科', '内分泌科', '肾内科']
DIAGNOSES = ['高血压病2级', '2型糖尿病', '慢性胃炎', '腰椎间盘突出症', '上呼吸道感染', '慢性阻塞性肺疾病']
ADVICES = ['注意休息，避免劳累', '低盐低脂饮食', '定期复查血糖', '继续口服药物', '加强康复锻炼', '1个月后门诊随访']

def random_patient():
    """生成随机患者信息字典"""
    name = fake.name()
    sex = random.choice(['男', '女'])
    age = random.randint(1, 90)
    hospital_id = str(random.randint(100000, 999999))
    department = random.choice(DEPARTMENTS)
    bed_num = f"{random.randint(1, 50)}床"
    adm = fake.date_between(start_date='-30d', end_date='-1d')
    dis = adm + timedelta(days=random.randint(1, 14))
    diagnosis = random.choice(DIAGNOSES)
    advice = random.choice(ADVICES)
    return {
        "姓名": name, "性别": sex, "年龄": age, "住院号": hospital_id,
        "科室": department, "床号": bed_num,
        "入院日期": adm.strftime('%Y-%m-%d'), "出院日期": dis.strftime('%Y-%m-%d'),
        "诊断": diagnosis, "医嘱": advice
    }

# 实体类型映射
FIELD2TYPE = {
    "姓名": "NAME", "性别": "SEX", "年龄": "AGE", "住院号": "HOSPITAL_NUM",
    "科室": "DEPARTMENT", "床号": "BED_NUM", "入院日期": "ADMISSION_DATE",
    "出院日期": "DISCHARGE_DATE", "诊断": "DIAGNOSIS", "医嘱": "ADVICE"
}

def build_text_and_entities(info, use_table):
    """
    生成全文文本和实体列表（位置信息准确）
    返回: (full_text, entities)
    """
    entities = []
    full_text = ""
    headers = list(info.keys())   # 统一获取字段顺序，保证稳定

    if use_table:
        # ---------- 表格模式 ----------
        # 先构造标题行和数值行（均使用制表符分隔）
        header_line = "\t".join(headers) + "\n"
        value_parts = [str(info[h]) for h in headers]
        value_line = "\t".join(value_parts) + "\n"

        full_text = header_line + value_line
        value_start = len(header_line)  # 数值行在全文中的起始偏移

        # 在数值行内逐个字段计算偏移（无需依赖find）
        current_pos = 0  # 数值行内的偏移（从0开始）
        for i, h in enumerate(headers):
            value = str(info[h])
            start = value_start + current_pos
            end = start + len(value)
            entities.append({"type": FIELD2TYPE[h], "start": start, "end": end})
            # 更新下一个字段的起始偏移：当前值长度 + 制表符（最后一个字段无制表符）
            current_pos += len(value) + (1 if i < len(headers)-1 else 0)

    else:
        # ---------- 纯文本模式 ----------
        # 1. 构造描述句，并精准记录其中字段的起止位置
        #    采用“零件拼接并累加长度”的方式，拒绝硬编码偏移
        desc_parts = []
        # -- 姓名 --
        name = info['姓名']
        start_name = len("患者")          # 开头“患者”的长度
        end_name = start_name + len(name)
        entities.append({"type": "NAME", "start": start_name, "end": end_name})
        desc_parts.append("患者")
        desc_parts.append(name)

        # -- 性别 --
        desc_parts.append("，")
        sex = info['性别']
        start_sex = sum(len(p) for p in desc_parts)
        end_sex = start_sex + len(sex)
        entities.append({"type": "SEX", "start": start_sex, "end": end_sex})
        desc_parts.append(sex)

        # -- 年龄 --
        desc_parts.append("，")
        age_str = str(info['年龄'])
        start_age = sum(len(p) for p in desc_parts)
        end_age = start_age + len(age_str)
        entities.append({"type": "AGE", "start": start_age, "end": end_age})
        desc_parts.append(age_str)

        # -- 诊断 --
        desc_parts.append("岁，因“")
        diag = info['诊断']
        start_diag = sum(len(p) for p in desc_parts)
        end_diag = start_diag + len(diag)
        entities.append({"type": "DIAGNOSIS", "start": start_diag, "end": end_diag})
        desc_parts.append(diag)
        desc_parts.append("”入院。\n")

        # 拼接描述句并加入全文
        desc_line = "".join(desc_parts)
        full_text += desc_line

        # 2. 逐行添加“字段：值”部分，并记录实体位置
        for h in headers:
            value = str(info[h])
            line = f"{h}：{value}\n"
            line_start = len(full_text)             # 本行在全文中的开始位置
            value_offset_in_line = len(f"{h}：")    # 值在本行内的偏移
            start = line_start + value_offset_in_line
            end = start + len(value)
            entities.append({"type": FIELD2TYPE[h], "start": start, "end": end})
            full_text += line

    return full_text, entities

def create_labels(text, entities):
    """根据实体列表生成BIO标签序列"""
    labels = ['O'] * len(text)
    for ent in entities:
        for i in range(ent['start'], ent['end']):
            if i == ent['start']:
                labels[i] = f'B-{ent["type"]}'
            else:
                labels[i] = f'I-{ent["type"]}'
    return labels


# -------------------- 主程序 --------------------
OUT_DIR = "generated"
os.makedirs(OUT_DIR, exist_ok=True)
train_file = "train_data.jsonl"

with open(train_file, 'w', encoding='utf-8') as f_train:
    for i in range(500):
        info = random_patient()
        use_table = random.choice([True, False])

        # 1. 生成 docx 文件
        doc = Document()
        if use_table:
            table = doc.add_table(rows=2, cols=len(info))
            headers = list(info.keys())
            for j, h in enumerate(headers):
                table.cell(0, j).text = h
            for j, h in enumerate(headers):
                table.cell(1, j).text = str(info[h])
        else:
            doc.add_paragraph(f"患者{info['姓名']}，{info['性别']}，{info['年龄']}岁，因“{info['诊断']}”入院。")
            for k, v in info.items():
                doc.add_paragraph(f"{k}：{v}")
        doc.save(os.path.join(OUT_DIR, f"patient_{i+1:04d}.docx"))

        # 2. 生成训练数据
        full_text, entities = build_text_and_entities(info, use_table)
        labels = create_labels(full_text, entities)
        record = {
            "file": f"patient_{i+1:04d}.docx",
            "text": full_text,
            "labels": labels
        }
        f_train.write(json.dumps(record, ensure_ascii=False) + "\n")

print(f"已生成 500 份 docx 至 {OUT_DIR}/")
print(f"训练数据已保存为 {train_file}")
